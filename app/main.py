"""Сервис очистки PII: POST /clean принимает текст или файл, возвращает очищенный текст.

Режимы (query-параметр mode):
  fast — только regex + NER (миллисекунды, без LLM);
  full — + LLM: контекстные даты и второй проход-аудит (по умолчанию).
"""
import io, time

from fastapi import FastAPI, UploadFile, File, Form, Query
from pydantic import BaseModel

from . import pii, llm

app = FastAPI(title="PII Filter",
              description="Очистка персональных данных из текста и документов перед отправкой во внешние LLM")


def extract_text(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if name.endswith(".docx"):
        import docx
        d = docx.Document(io.BytesIO(content))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:  # в таблицах договоров часто лежат реквизиты
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    # txt и всё остальное — как текст
    for enc in ("utf-8", "windows-1251", "cp866"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


class CleanTextIn(BaseModel):
    text: str


async def _run_clean(text: str, mode: str) -> dict:
    t0 = time.perf_counter()
    llm.usage.update({"llm_calls": 0, "prompt_tokens": 0, "completion_tokens": 0})
    use_llm = mode == "full" and not llm.MOCK_MODE
    result = await pii.clean_text(text, llm_classify_dates=llm.classify_dates if use_llm else None)
    if use_llm:
        extra = await llm.audit_residual_pii(result["cleaned_text"])
        result["cleaned_text"], result["removed"] = pii.apply_extra_removals(
            result["cleaned_text"], result["removed"], extra)
        result["entities_found"] = len(result["removed"])
    result["mode"] = mode if use_llm or mode == "fast" else "fast (MOCK_MODE)"
    result["took_ms"] = round((time.perf_counter() - t0) * 1000, 1)
    result["llm_usage"] = dict(llm.usage)
    return result


@app.post("/clean")
async def clean(payload: CleanTextIn, mode: str = Query("full", enum=["fast", "full"])):
    """Очистка сырого текста."""
    return await _run_clean(payload.text, mode)


@app.post("/clean/file")
async def clean_file(file: UploadFile = File(...), mode: str = Form("full")):
    """Очистка документа (PDF/DOCX/TXT): текст извлекается библиотеками, файл никуда не уходит."""
    text = extract_text(file.filename, await file.read())
    result = await _run_clean(text, mode)
    result["source_file"] = file.filename
    result["extracted_chars"] = len(text)
    return result


class RestoreIn(BaseModel):
    text: str          # ответ внешней LLM, содержащий плейсхолдеры
    removed: list[dict]  # карта подстановок — поле `removed` из ответа /clean


@app.post("/restore")
def restore(payload: RestoreIn):
    """Обратная подстановка: возвращает в ответ модели исходные значения.

    Полный цикл гейтвея: /clean -> внешняя LLM -> /restore. Карта подстановок
    хранится на стороне гейтвея и наружу не уходит.
    """
    return pii.restore_text(payload.text, payload.removed)


@app.get("/health")
def health():
    return {"ok": True, "mock_mode": llm.MOCK_MODE, "llm_model": llm.LLM_MODEL}
